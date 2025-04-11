import os
import json
import hashlib
from shutil import copyfile
from docx import Document
from tkinter import Tk, filedialog
import sys

def ShowProgressBar(Current, Total, Prefix="", Length=40):
    Percent = int((Current / Total) * 100)
    Filled = int(Length * Percent // 100)
    Bar = "█" * Filled + '-' * (Length - Filled)
    sys.stdout.write(f"\r{Prefix} |{Bar}| {Percent}% ({Current}/{Total})")
    sys.stdout.flush()

def Pause():
    input("\n按任意键结束...")

def CountImagesInDoc(Doc):
    Count = 0
    for Rel in Doc.part._rels.values():
        if "image" in Rel.target_ref:
            Count += 1
    return Count

def CountQuestionsInDoc(Doc):
    return sum(1 for Para in Doc.paragraphs if Para.text.strip().startswith("S") and "." in Para.text)

def BuildImageHashMap(Doc):
    ImageMap = {}
    for Rel in Doc.part._rels.values():
        if "image" in Rel.target_ref:
            Data = Rel.target_part.blob
            HashVal = hashlib.md5(Data).hexdigest()
            ImageMap[HashVal] = Data
    return ImageMap

def SaveQuestionImage(ImageMap, ImageDir, QuestionId, ImageIndex):
    ImageFolder = os.path.join(ImageDir, QuestionId)
    os.makedirs(ImageFolder, exist_ok=True)

    if not ImageMap:
        return ""

    HashVal, Data = list(ImageMap.items())[0]
    ImageName = f"{ImageIndex}.png"
    DstPath = os.path.join(ImageFolder, ImageName)
    with open(DstPath, "wb") as F:
        F.write(Data)
    ImageMap.pop(HashVal)
    return ImageName

def ExtractQuestions(Doc, ImageDir, OutputPath):
    ImageMap = BuildImageHashMap(Doc)
    Questions = []
    Current = None
    QuestionIndex = 0

    Total = len(Doc.paragraphs)
    print("\n📖 题目解析进度：")

    for I, Para in enumerate(Doc.paragraphs):
        Text = Para.text.strip()
        ShowProgressBar(I + 1, Total, Prefix="解析中")

        if not Text:
            continue

        if Text.startswith("S") and "." in Text:
            Current = {
                "题目ID": Text,
                "题目": {"文本": "", "图片": ""},
                "题目类型": "单选",
                "选项": [],
                "题目解析": "",
                "解析库": [],
            }
            Section = None
            ImageIndex = 0
            continue

        if Text in {"题目", "选项", "答案", "题目解析"}:
            Section = Text
            continue

        if not Current:
            continue

        if Section == "题目":
            Current["题目"]["文本"] += Text + " "
            if Current:
                Current["题目"]["图片"] = SaveQuestionImage(ImageMap, ImageDir, Current["题目ID"], ImageIndex)
                Questions.append(Current)
                QuestionIndex += 1
        elif Section == "选项" and Text[:1] in "ABCD":
            OptionText = Text[1:].strip()
            ImageIndex += 1
            OptionImage = SaveQuestionImage(ImageMap, ImageDir, Current["题目ID"], ImageIndex)
            Current["选项"].append({
                "文本": OptionText,
                "图片": OptionImage,
                "是否正确": False,
                "解析": ""
            })
        elif Section == "答案":
            AnswerLetter = Text.strip()
            for Index, Letter in enumerate("ABCD"):
                if Index < len(Current["选项"]):
                    Current["选项"][Index]["是否正确"] = (Letter == AnswerLetter)
        elif Section == "题目解析":
            Current["题目解析"] += Text + " "

    if Current:
        Current["题目"]["图片"] = SaveQuestionImage(ImageMap, ImageDir, Current["题目ID"], 0)
        Questions.append(Current)

    print("\n📦 图片保存进度：")
    for I in range(len(Questions)):
        ShowProgressBar(I + 1, len(Questions), Prefix="保存中")

    with open(OutputPath, "w", encoding="utf-8") as F:
        json.dump({
            "题库": Questions,
            "公共解析库": []
        }, F, ensure_ascii=False, indent=2)

    print("\n\n✅ 导出完成：", OutputPath)
    Pause()

def Main():
    try:
        print("📂 正在打开文件选择窗口...")
        Root = Tk()
        Root.withdraw()
        FilePath = filedialog.askopenfilename(title="选择 Word 文件", filetypes=[("Word 文件", "*.docx")])
        if not FilePath:
            print("❌ 未选择文件，程序结束。")
            Pause()
            return

        Doc = Document(FilePath)
        ImageCount = CountImagesInDoc(Doc)
        QuestionCount = CountQuestionsInDoc(Doc)

        print(f"📷 Word 中嵌入图片数量：{ImageCount}")
        print(f"🧾 检测到题目数量：{QuestionCount}")
        print("🚀 准备开始导入...\n")

        ExtractQuestions(Doc=Doc, ImageDir="Images", OutputPath="输出题库.json")

    except Exception as E:
        print("\n❌ 程序出错：", E)
        Pause()

if __name__ == "__main__":
    Main()
